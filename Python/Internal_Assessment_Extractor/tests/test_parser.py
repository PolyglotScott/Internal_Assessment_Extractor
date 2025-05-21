"""
test_parser.py

Unit tests for the parse_document function using in-memory mock Word documents
created with python-docx. Covers valid, malformed, and complex Q&A scenarios.
"""

import unittest
import tempfile
import os
from docx import Document
from src.parser import parse_document

class TestParseDocument(unittest.TestCase):

    def create_temp_docx(self, paragraphs=None, tables=None):
        """
        Helper to create and save a temporary docx file.

        Args:
            paragraphs (list): List of (text, heading_level) tuples
            tables (list): List of lists, each a table with rows of cell texts

        Returns:
            str: Path to temporary docx file
        """
        doc = Document()
        if paragraphs:
            for text, level in paragraphs:
                doc.add_heading(text, level=level)

        if tables:
            for table_rows in tables:
                table = doc.add_table(rows=len(table_rows), cols=1)
                for i, row_text in enumerate(table_rows):
                    table.cell(i, 0).text = row_text

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)
        return tmp.name

    def tearDown(self):
        if hasattr(self, 'doc_path') and os.path.exists(self.doc_path):
            os.remove(self.doc_path)

    def test_valid_single_qa(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("Chapter 1", 1), ("Process A", 2)],
            tables=[["CONCEPT CHECK\nASK participants: What is 2 + 2? ANSWER: 4"]]
        )
        df = parse_document(self.doc_path)
        self.assertEqual(len(df), 1)
        self.assertEqual(df.loc[0, 'Questions'], "What is 2 + 2?")
        self.assertEqual(df.loc[0, 'Answer'], "4")

    def test_multiple_qa_in_table(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("Chapter 2", 1), ("Process B", 2)],
            tables=[[
                "CONCEPT CHECK\nASK participants: First Q? ANSWER: First A",
                "CONCEPT CHECK\nASK participants: Second Q? ANSWER: Second A"
            ]]
        )
        df = parse_document(self.doc_path)
        self.assertEqual(len(df), 2)
        self.assertIn("First Q?", df['Questions'].values)
        self.assertIn("Second A", df['Answer'].values)

    def test_malformed_question_only(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("Chapter 3", 1), ("Process C", 2)],
            tables=[["CONCEPT CHECK\nASK participants: What is broken?"]]
        )
        df = parse_document(self.doc_path)
        self.assertTrue(df.empty)

    def test_malformed_answer_only(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("Chapter 3", 1), ("Process D", 2)],
            tables=[["CONCEPT CHECK\nANSWER: Missing question"]]
        )
        df = parse_document(self.doc_path)
        self.assertTrue(df.empty)

    def test_no_concept_check(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("Chapter 4", 1), ("Process E", 2)],
            tables=[["Random text without format"]]
        )
        df = parse_document(self.doc_path)
        self.assertTrue(df.empty)

    def test_multiple_tables_with_mixed_qas(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("Chapter 5", 1), ("Process F", 2)],
            tables=[
                ["CONCEPT CHECK\nASK participants: Q1? ANSWER: A1"],
                ["Some junk", "CONCEPT CHECK\nASK participants: Q2? ANSWER: A2"]
            ]
        )
        df = parse_document(self.doc_path)
        self.assertEqual(len(df), 2)
        self.assertIn("Q1?", df['Questions'].values)
        self.assertIn("A2", df['Answer'].values)

    def test_preserves_chapter_process(self):
        self.doc_path = self.create_temp_docx(
            paragraphs=[("My Chapter", 1), ("My Process", 2)],
            tables=[["CONCEPT CHECK\nASK participants: Check me? ANSWER: Checked!"]]
        )
        df = parse_document(self.doc_path)
        self.assertEqual(df.loc[0, 'Chapter'], "My Chapter")
        self.assertEqual(df.loc[0, 'Process'], "My Process")

if __name__ == "__main__":
    unittest.main()
