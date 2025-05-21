from docx import Document
import pandas as pd

# Load the document
doc_path = doc_path = "C:/Users/vivif/Documents/Code/Python/Internal_Assessment_Extractor/input/Module 3 - NTS Financial Services Systems Overview - Facilitator Guide - V1.docx"
doc = Document(doc_path)

# Variables to track current chapter and process
current_chapter = ""
current_process = ""

# Results to build table
table_data = []

# Helper function to extract question and answer from cell text
def extract_question_answer(cell_text):
    question = ""
    answer = ""
    if 'ASK participants:' in cell_text and 'ANSWER:' in cell_text:
        try:
            question = cell_text.split('ASK participants:')[1].split('ANSWER:')[0].strip()
            answer = cell_text.split('ANSWER:')[1].strip()
        except IndexError:
            pass
    return question, answer

# Iterate through paragraphs and tables using styling
for para in doc.paragraphs:
    style = para.style.name
    if style in ['Heading 1', 'Header 2']:
        current_chapter = para.text.strip()
        table_data.append([current_chapter, '', ''])
    elif style in ['Heading 2', 'Header 3']:
        current_process = para.text.strip()
        table_data.append([current_process, '', ''])

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if 'CONCEPT CHECK' in cell.text and 'ASK participants:' in cell.text and 'ANSWER:' in cell.text:
                question, answer = extract_question_answer(cell.text)
                if question and answer:
                    table_data.append([question, answer, '/1'])

# Convert to DataFrame for formatting and display
df = pd.DataFrame(table_data, columns=['Questions', 'Answer', 'Marks'])
df.head(15)  # Preview first 15 rows