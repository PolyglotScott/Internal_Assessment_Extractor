"""
doc_parser.py

Parses DOCX documents to extract chapters, processes, subsections, and question-answer pairs
for internal assessment extraction, using both numbering patterns and heading styles.
"""

import logging
import re
import multiprocessing
import sys
from docx import Document
import pandas as pd

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")


def clean(text):
    """
    Clean and normalize whitespace in a text string.
    """
    return ' '.join(text.strip().split()) if text else ''


def extract_structure(para, current):
    """
    Update the current structure dictionary based on paragraph style and text.
    """
    style = para.style.name.lower()
    text = clean(para.text)
    if style.startswith("heading 1") and text:
        current['chapter'] = text
        current['chapter_style'] = style
    elif style.startswith("heading 2") and text:
        current['process'] = text
        current['process_style'] = style
    elif style.startswith("heading 3") and text:
        current['subsection'] = text
        current['subsection_style'] = style
    return current


def extract_question_answer(cell_text):
    """
    Extract question, answer, and question type from a cell containing both.
    """
    question_type = None
    question = None
    answer = None

    # Identify question type
    if "CHALLENGES CONCEPT CHECK" in cell_text.upper():
        question_type = "Challenges Concept Check Question"

    # Extract question and answer using regex
    match = re.search(
        r"(ASK participants:.*?)(ANSWER:.*)",
        cell_text,
        re.IGNORECASE | re.DOTALL
    )
    if match:
        question = match.group(1).replace("ASK participants:", "").strip()
        answer = match.group(2).replace("ANSWER:", "").strip()
    else:
        # Fallback: try to find just the answer
        answer_match = re.search(r"ANSWER:(.*)", cell_text, re.IGNORECASE | re.DOTALL)
        if answer_match:
            answer = answer_match.group(1).strip()
        # Try to find just the question
        question_match = re.search(r"ASK participants:(.*)", cell_text, re.IGNORECASE | re.DOTALL)
        if question_match:
            question = question_match.group(1).strip()

    return question, answer, question_type


def extract_qa_from_cell(cell):
    """
    Extract all Q&A pairs from a table cell.
    """
    qas = []
    last_concept_check = None
    paragraphs = cell.paragraphs
    for idx, para in enumerate(paragraphs):
        style = para.style.name.lower()
        text = clean(para.text)
        if "concept check" in text.lower() and style in ["normal", "body text"]:
            last_concept_check = "CONCEPT CHECK"
        if "ask participants:" in text.lower():
            question = text
            answer = ""
            # Try to find answer in the same or next paragraph
            if "answer:" in text.lower():
                parts = text.split("answer:", 1)
                question = parts[0].strip()
                answer = parts[1].strip() if len(parts) > 1 else ""
            elif idx + 1 < len(paragraphs):
                next_text = clean(paragraphs[idx + 1].text)
                if next_text.lower().startswith("answer:"):
                    answer = next_text[7:].strip()
            qas.append((last_concept_check, question, answer))
    return qas


def parse_document(path):
    """
    Parse a DOCX document and extract structured Q&A data as a DataFrame.
    """
    doc = Document(path)
    current = {
        "chapter": "", "chapter_style": "",
        "process": "", "process_style": "",
        "subsection": "", "subsection_style": ""
    }
    data = []
    seen_questions = set()  # Track unique questions

    for para in doc.paragraphs:
        current = extract_structure(para, current)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Update structure if headings are in table cells
                for para in cell.paragraphs:
                    current = extract_structure(para, current)
                # Extract all Q&A pairs from this cell
                qas = extract_qa_from_cell(cell)
                for qtype, question, answer in qas:
                    # Only add if question is not already seen
                    question_key = question.strip().lower() if question else ""
                    if question_key and question_key not in seen_questions:
                        seen_questions.add(question_key)
                        data.append({
                            "QuestionType": qtype or "Unknown",
                            "Questions": question,
                            "Answer": answer,
                            "Marks": "/1",
                            "Chapter": current['chapter'],
                            "ChapterStyle": current['chapter_style'],
                            "Process": current['process'],
                            "ProcessStyle": current['process_style'],
                            "Subsection": current['subsection'],
                            "SubsectionStyle": current['subsection_style']
                        })

    df = pd.DataFrame(data)
    logging.info("Extracted %d unique Q&A pairs.", len(df))
    return df


def parse_documents_in_parallel(doc_paths):
    """
    Parse multiple DOCX documents in parallel and return a list of DataFrames.
    """
    with multiprocessing.Pool() as pool:
        dfs = pool.map(parse_document, doc_paths)
    return dfs


if __name__ == "__main__":
    import argparse
    import os
    from pathlib import Path

    # Always resolve input path relative to this script's parent directory
    script_dir = Path(__file__).resolve().parent.parent
    input_dir = script_dir / "input"
    input_dir.mkdir(exist_ok=True)

    # Find all .docx files in the input directory
    docx_files = list(input_dir.glob("*.docx"))

    parser = argparse.ArgumentParser(
        description=(
            "Parse DOCX documents for internal assessment extraction using doc_parser.py."
        )
    )
    parser.add_argument(
        "--input",
        type=str,
        default=None,
        help=(
            "Path to the DOCX file. If not provided, the first .docx file in the input "
            "folder will be used."
        )
    )
    args = parser.parse_args()

    # Determine which file to use
    if args.input:
        selected_doc_path = os.path.abspath(args.input)
    elif docx_files:
        selected_doc_path = str(docx_files[0])  # pylint: disable=invalid-name
        print(f"No --input provided. Using: {selected_doc_path}")
    else:
        logging.error("No DOCX files found in %s", input_dir)
        print(f"Error: No DOCX files found in the input directory: {input_dir}")
        sys.exit(1)

    if not os.path.isfile(selected_doc_path):
        logging.error("Input file not found: %s", selected_doc_path)
        print(f"Error: Input file not found: {selected_doc_path}")
        sys.exit(1)

    try:
        df_questions = parse_document(selected_doc_path)
        print("DataFrame Summary:")
        print(df_questions.info())
        print(df_questions.head(15))
    except (OSError, ValueError) as exc:
        logging.error("An error occurred during parsing: %s", exc)
        print(f"An error occurred during parsing: {exc}")
