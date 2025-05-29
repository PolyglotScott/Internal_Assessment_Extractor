"""
injector.py

Provides a function to inject cleaned DataFrame data into a document template.
This basic implementation locates a cell in a DOCX table by row and column index and
modifies its value.
"""

from pathlib import Path
from docx import Document

def inject_data_into_table_cell(
    template_name: str,
    row_idx: int,
    col_idx: int,
    new_value: str,
    output_path: str
) -> None:
    """
    Injects a value into a specific cell of the first table in the template DOCX file.

    Args:
        template_name (str): Name of the template file to use (must be in 'templates' directory).
        row_idx (int): Row index of the cell to modify (0-based).
        col_idx (int): Column index of the cell to modify (0-based).
        new_value (str): The value to insert into the cell.
        output_path (str): Path to save the modified document.
    """
    template_dir = Path(__file__).resolve().parent.parent / "templates"
    template_path = template_dir / template_name

    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found: {template_path}")

    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("No tables found in the template document.")

    table = doc.tables[0]
    if row_idx >= len(table.rows) or col_idx >= len(table.columns):
        raise IndexError("Specified cell is out of table bounds.")

    cell = table.cell(row_idx, col_idx)
    cell.text = new_value

    doc.save(output_path)

# Example usage (uncomment to use directly):
# inject_data_into_table_cell(
#     template_name="assessment_template.docx",
#     row_idx=1,
#     col_idx=2,
#     new_value="Injected Value",
#     output_path="output/injected_result.docx"
# )
