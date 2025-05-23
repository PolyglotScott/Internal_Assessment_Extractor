"""
parser.py

Responsible for parsing DOCX documents to extract chapters, processes,
and question-answer pairs for internal assessment extraction.

This module focuses on using heading styles for structure, with optional
cross-checking against the Table of Contents (TOC) if present.
"""

import logging
from docx import Document
import pandas as pd

logger = logging.getLogger(__name__)


def extract_chapters_processes_by_headings(doc):
    """
    Extract chapters and processes using heading styles.
    Returns a list of chapters and a mapping of chapters to processes.
    """
    chapters = []
    chapter_process_map = {}
    current_chapter = None

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if style == "Heading 1":
            current_chapter = text
            chapters.append(current_chapter)
            chapter_process_map[current_chapter] = []
        elif style == "Heading 2" and current_chapter:
            chapter_process_map[current_chapter].append(text)
    return chapters, chapter_process_map


def extract_toc_structure(doc):
    """
    Extracts the Table of Contents structure from the document.
    Returns a list of (level, text) tuples.
    """
    toc = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if style.startswith("TOC "):
            try:
                level = int(style.split("TOC ")[1])
                toc.append((level, text))
            except Exception:
                continue
    return toc


def crosscheck_headings_with_toc(chapters, chapter_process_map, toc):
    """
    Optionally cross-checks heading-based structure with TOC.
    Logs discrepancies for review.
    """
    toc_chapters = [text for level, text in toc if level == 1]
    toc_processes = [text for level, text in toc if level == 2]
    missing_in_headings = set(toc_chapters) - set(chapters)
    missing_in_toc = set(chapters) - set(toc_chapters)
    if missing_in_headings:
        logger.warning("Chapters in TOC but not in headings: %s", missing_in_headings)
    if missing_in_toc:
        logger.warning("Chapters in headings but not in TOC: %s", missing_in_toc)
    # Optionally, cross-check processes as well


def extract_qa_from_tables(doc, chapter, process):
    """
    Extract question-answer pairs from tables in the document.
    Returns a list of [question, answer, marks, chapter, process].
    """
    table_data = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if 'CONCEPT CHECK' in cell_text and 'ASK participants:' in cell_text and 'ANSWER:' in cell_text:
                    try:
                        question = cell_text.split('ASK participants:')[1].split('ANSWER:')[0].strip()
                        answer = cell_text.split('ANSWER:')[1].strip()
                        table_data.append([question, answer, '/1', chapter, process])
                    except Exception:
                        continue
    return table_data


def parse_document(doc_path: str):
    """
    Parses the DOCX document, extracting chapters, processes, and Q&A.

    Args:
        doc_path (str): Path to the DOCX file.

    Returns:
        pd.DataFrame: DataFrame containing Questions, Answers, Marks, Chapter, and Process.
    """
    try:
        doc = Document(doc_path)
    except Exception as exc:
        logger.error("Failed to load document: %s | Error: %s. Please check if the file path is correct and the file is a valid DOCX document.", doc_path, exc)
        raise

    # 1. Parse using heading styles
    chapters, chapter_process_map = extract_chapters_processes_by_headings(doc)
    logger.info("Chapters found by headings: %d", len(chapters))
    for chapter in chapters:
        logger.info("Chapter: '%s' | Processes: %d", chapter, len(chapter_process_map[chapter]))
        for process in chapter_process_map[chapter]:
            logger.info("  Process: '%s'", process)

    # 2. Optionally cross-check with TOC
    toc = extract_toc_structure(doc)
    if toc:
        crosscheck_headings_with_toc(chapters, chapter_process_map, toc)

    # 3. Extract Q&A from tables for each chapter and process
    table_data = []
    for chapter in chapters:
        for process in chapter_process_map[chapter]:
            table_data.extend(extract_qa_from_tables(doc, chapter, process))

    df = pd.DataFrame(table_data, columns=['Questions', 'Answer', 'Marks', 'Chapter', 'Process'])
    df_cleaned = df[(df['Questions'] != '') | (df['Answer'] != '')].reset_index(drop=True)
    return df_cleaned


if __name__ == "__main__":
    import argparse
    import os
    from pathlib import Path

    # Always resolve input path relative to this script's parent directory
    script_dir = Path(__file__).resolve().parent.parent
    input_dir = script_dir / "input"
    input_dir.mkdir(exist_ok=True)

    # Find all .docx files in the input directory
    docx_files = list(input_dir.glob("*.docx"))

    parser = argparse.ArgumentParser(description="Parse DOCX documents for internal assessment extraction.")
    parser.add_argument(
        "--input",
        type=str,
        default=None,
        help="Path to the DOCX file. If not provided, the first .docx file in the input folder will be used."
    )
    args = parser.parse_args()

    # Determine which file to use
    if args.input:
        doc_path = os.path.abspath(args.input)
    elif docx_files:
        doc_path = str(docx_files[0])
        print(f"No --input provided. Using: {doc_path}")
    else:
        logger.error("No DOCX files found in the input directory: %s", input_dir)
        print(f"Error: No DOCX files found in the input directory: {input_dir}")
        exit(1)

    if not os.path.isfile(doc_path):
        logger.error("Input file not found: %s", doc_path)
        print(f"Error: Input file not found: {doc_path}")
        exit(1)

    try:
        df_questions = parse_document(doc_path)
        print("DataFrame Summary:")
        print(df_questions.info())
        print(df_questions.head(15))
    except Exception as exc:
        logger.error("An error occurred during parsing: %s", exc)
        print(f"An error occurred during parsing: {exc}")
        exit(2)
